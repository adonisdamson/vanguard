#!/usr/bin/env python3
"""
Generate a Microsoft Word document with all operator accounts and their credentials.
Outputs: NDC_Vanguard_Credentials.docx on the Desktop.
"""
import json, urllib.request, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SUPABASE_URL = "https://gnolfngnwyuwubqzopzl.supabase.co"
SERVICE_ROLE_KEY = (
    "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9."
    "eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Imdub2xmbmdud3l1d3VicXpvcHpsIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc4MzEwMjU3OCwiZXhwIjoyMDk4Njc4NTc4fQ."
    "FVNU0eby__ODH5WjIKz-QSwqKJIUoGe8oKFx-pQwLxk"
)
PASSWORD = "temawestndc@2026!"

# NDC colors
NDC_GREEN  = RGBColor(0x00, 0x6B, 0x3F)
NDC_RED    = RGBColor(0xCE, 0x11, 0x26)
NDC_BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
NDC_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF4, 0xF6, 0xF5)
MID_GREY   = RGBColor(0xD8, 0xDE, 0xE2)

ROLE_LABELS = {
    "admin": "System Administrator",
    "higher_authority": "Coordinator",
    "manager": "Manager",
    "personnel": "Field Officer",
}

def fetch_users():
    """Fetch all app_users ordered by role then name."""
    url = (
        f"{SUPABASE_URL}/rest/v1/app_users"
        "?select=id,full_name,email,phone,role,is_active"
        "&order=role.asc,full_name.asc"
        "&limit=1000"
    )
    req = urllib.request.Request(url, headers={
        "apikey": SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SERVICE_ROLE_KEY}",
        "Prefer": "return=representation",
    })
    with urllib.request.urlopen(req) as resp:
        return json.loads(resp.read())

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="D8DEE2"):
    """Add thin borders to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading_paragraph(doc, text, size=14, bold=True, color=NDC_GREEN, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def main():
    print("Fetching operator accounts...")
    users = fetch_users()
    print(f"  {len(users)} accounts found.")

    # Group by role
    groups = {}
    role_order = ["admin", "higher_authority", "manager", "personnel"]
    for role in role_order:
        groups[role] = [u for u in users if u["role"] == role]

    doc = Document()

    # Page margins — tight to fit more per page
    section = doc.sections[0]
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── Title block ──────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("VANGUARD — NDC MEMBERSHIP SYSTEM")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = NDC_GREEN

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Tema West Constituency — Operator Login Credentials")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = NDC_BLACK

    sub_p2 = doc.add_paragraph()
    sub_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run2 = sub_p2.add_run(f"Confidential — {len(users)} accounts | Default password: {PASSWORD}")
    sub_run2.font.size = Pt(10)
    sub_run2.font.color.rgb = RGBColor(0x5B, 0x66, 0x70)
    sub_run2.italic = True

    doc.add_paragraph()  # spacer

    # ── Per-role tables ───────────────────────────────────────────────────────────
    col_widths = [Cm(6.0), Cm(4.5), Cm(3.8), Cm(4.5)]
    headers = ["Full Name", "Login (Phone)", "Role", "Password"]

    for role in role_order:
        batch = groups[role]
        if not batch:
            continue

        role_label = ROLE_LABELS.get(role, role.title())
        add_heading_paragraph(
            doc, f"{role_label}s  ({len(batch)})",
            size=12, color=NDC_GREEN, space_before=8, space_after=4
        )

        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'Table Grid'

        # Header row
        hdr = table.rows[0]
        for j, (htext, w) in enumerate(zip(headers, col_widths)):
            cell = hdr.cells[j]
            cell.width = w
            set_cell_bg(cell, "006B3F")
            set_cell_borders(cell, "006B3F")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(htext)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = NDC_WHITE

        # Data rows
        for idx, u in enumerate(batch):
            phone = u.get("phone") or u.get("email", "").split("@")[0]
            row = table.add_row()
            bg = "F4F6F5" if idx % 2 == 0 else "FFFFFF"

            values = [
                u.get("full_name", ""),
                phone,
                ROLE_LABELS.get(u["role"], u["role"]),
                PASSWORD,
            ]
            for j, (val, w) in enumerate(zip(values, col_widths)):
                cell = row.cells[j]
                cell.width = w
                set_cell_bg(cell, bg)
                set_cell_borders(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val)
                run.font.size = Pt(9)
                run.font.color.rgb = NDC_BLACK
                if j == 3:  # password column — monospace feel
                    run.bold = True

        doc.add_paragraph()  # spacer between sections

    # ── Instructions ─────────────────────────────────────────────────────────────
    doc.add_page_break()

    add_heading_paragraph(doc, "Login Instructions", size=13, color=NDC_GREEN, space_before=0)

    instructions = [
        ("Mobile App (Android APK)", [
            "Download and install the Vanguard APK from the party office.",
            "Open the app and tap 'Sign In'.",
            "Enter your phone number in the Login field (numbers only, e.g. 0244123456).",
            f"Enter the password: {PASSWORD}",
            "Tap Sign In. You will be taken to your role-specific dashboard.",
        ]),
        ("Web App", [
            "Visit the web app URL provided by your constituency coordinator.",
            "Use the same phone number and password as above.",
        ]),
        ("Changing Your Password", [
            "After first login, go to your Profile (top-right avatar).",
            "Tap 'Change password' and follow the prompts.",
            "Choose a strong personal password and keep it confidential.",
        ]),
        ("Support", [
            "Contact your System Administrator if you cannot log in.",
            "Do not share your password with anyone.",
            "This document is CONFIDENTIAL — do not distribute beyond authorised personnel.",
        ]),
    ]

    for title, points in instructions:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = NDC_BLACK

        for point in points:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after = Pt(1)
            run = bp.add_run(point)
            run.font.size = Pt(10)
            run.font.color.rgb = NDC_BLACK

    # ── Footer note ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        "National Democratic Congress · Tema West Constituency · Vanguard System · CONFIDENTIAL"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x9B, 0xA1, 0xA8)
    footer_run.italic = True

    out_path = os.path.expanduser("~/Desktop/NDC_Vanguard_Credentials.docx")
    doc.save(out_path)
    print(f"\nDocument saved: {out_path}")
    print(f"Contains {len(users)} operator accounts across {len([r for r in role_order if groups[r]])} role groups.")

if __name__ == "__main__":
    main()
